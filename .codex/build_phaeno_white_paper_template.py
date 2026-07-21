from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(r"D:\__dev\phaeno-portal")
OUTPUT = ROOT / "docs" / "templates" / "Phaeno-White-Paper-Template.docx"
WORK = ROOT / ".codex" / "white-paper-template-work"
LOGO = ROOT / "website" / "public" / "images" / "phaeno124x40.webp"
COVER_IMAGE = ROOT / "website" / "public" / "images" / "molecules.png"


# Phaeno Website design-system colors.
RNA_50 = "EDF7FA"
RNA_100 = "DCEEF4"
RNA_600 = "156082"
RNA_900 = "0E2841"
GREEN_50 = "F6F8F2"
GREEN_100 = "E9EFEA"
GREEN_500 = "789946"
GREEN_700 = "526832"
AMBER_100 = "FFF7E8"
AMBER_400 = "FEC950"
ORANGE_600 = "B35D0C"
WHITE = "FFFFFF"
NEUTRAL_50 = "F7F8F8"
NEUTRAL_200 = "D9D9D9"
NEUTRAL_300 = "C9C9C9"
NEUTRAL_500 = "7F7F7F"
NEUTRAL_600 = "595959"
NEUTRAL_900 = "1D1D1D"

FONT = "Segoe UI"  # Inter's first installed fallback in the Website design system.
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def ensure_child(parent, tag: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def remove_children(parent, tag: str) -> None:
    for child in list(parent.findall(qn(tag))):
        parent.remove(child)


def set_run_font(run, *, size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None,
                 all_caps: bool | None = None, tracking_twips: int | None = None) -> None:
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attr), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if all_caps is not None:
        run.font.all_caps = all_caps
    if tracking_twips is not None:
        spacing = rpr.find(qn("w:spacing"))
        if spacing is None:
            spacing = OxmlElement("w:spacing")
            rpr.append(spacing)
        spacing.set(qn("w:val"), str(tracking_twips))


def set_style_font(style, *, size: float, color: str, bold: bool = False,
                   italic: bool = False) -> None:
    style.font.name = FONT
    style.font.size = Pt(size)
    style.font.color.rgb = rgb(color)
    style.font.bold = bold
    style.font.italic = italic
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attr), FONT)


def set_keep_with_next(style, value: bool = True) -> None:
    style.paragraph_format.keep_with_next = value


def set_widow_control(paragraph_or_style, value: bool = True) -> None:
    ppr = paragraph_or_style.element.get_or_add_pPr() if hasattr(paragraph_or_style, "element") else paragraph_or_style._p.get_or_add_pPr()
    widow = ppr.find(qn("w:widowControl"))
    if widow is None:
        widow = OxmlElement("w:widowControl")
        ppr.append(widow)
    widow.set(qn("w:val"), "1" if value else "0")


def set_paragraph_shading(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_paragraph_left_border(paragraph, color: str, *, size: int = 18, space: int = 8) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    left = pbdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        pbdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins: dict[str, int]) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    tcmar = ensure_child(tcpr, "w:tcMar")
    for side in ("top", "bottom", "start", "end"):
        margin = ensure_child(tcmar, f"w:{side}")
        margin.set(qn("w:w"), str(margins[side]))
        margin.set(qn("w:type"), "dxa")


def apply_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblpr = table._tbl.tblPr
    tblw = ensure_child(tblpr, "w:tblW")
    tblw.set(qn("w:type"), "dxa")
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tblind = ensure_child(tblpr, "w:tblInd")
    tblind.set(qn("w:type"), "dxa")
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    layout = ensure_child(tblpr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for col_index, width in enumerate(widths_dxa):
        table.columns[col_index].width = Twips(width)
    for row in table.rows:
        row.height = None
        for col_index, cell in enumerate(row.cells):
            width = widths_dxa[col_index]
            cell.width = Twips(width)
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = ensure_child(tcpr, "w:tcW")
            tcw.set(qn("w:type"), "dxa")
            tcw.set(qn("w:w"), str(width))
            set_cell_margins(cell, CELL_MARGINS_DXA)


def set_table_borders(table, color: str = NEUTRAL_200) -> None:
    tblpr = table._tbl.tblPr
    borders = ensure_child(tblpr, "w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = ensure_child(borders, f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    trpr = row._tr.get_or_add_trPr()
    header = trpr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        trpr.append(header)
    header.set(qn("w:val"), "true")


def style_table(table, widths_dxa: list[int], *, emphasize_column: int | None = None) -> None:
    apply_table_geometry(table, widths_dxa)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, RNA_50)
            elif emphasize_column is not None and col_index == emphasize_column:
                set_cell_shading(cell, GREEN_50)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_run_font(
                        run,
                        size=8.5,
                        color=RNA_900 if row_index == 0 else NEUTRAL_900,
                        bold=row_index == 0,
                    )


def add_field(paragraph, instruction: str, display: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    instruction_run._r.append(instr)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    display_run = paragraph.add_run(display)
    set_run_font(display_run, size=8.5, color=NEUTRAL_600)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_hyperlink(paragraph, text: str, anchor: str):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attr), FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), NEUTRAL_900)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    rpr.extend([fonts, color, underline])
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.extend([rpr, text_el])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    insert_at = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_at, start)
    paragraph._p.append(end)


def set_picture_alt_text(inline_shape, description: str) -> None:
    inline_shape._inline.docPr.set("descr", description)


def add_spacer(doc: Document, points: float) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(points)
    paragraph.paragraph_format.line_spacing = Pt(1)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def add_callout(doc: Document, label: str, text: str, *, kind: str = "takeaway"):
    paragraph = doc.add_paragraph(style="Phaeno Callout")
    if kind == "evidence":
        set_paragraph_shading(paragraph, AMBER_100)
        set_paragraph_left_border(paragraph, AMBER_400)
    else:
        set_paragraph_shading(paragraph, GREEN_50)
        set_paragraph_left_border(paragraph, GREEN_500)
    lead = paragraph.add_run(f"{label}  ")
    set_run_font(lead, size=9.5, color=RNA_900, bold=True)
    body = paragraph.add_run(text)
    set_run_font(body, size=9.5, color=NEUTRAL_900)
    return paragraph


def add_heading(doc: Document, text: str, level: int, bookmark: str, bookmark_id: int,
                *, page_break_before: bool = False):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.page_break_before = page_break_before
    add_bookmark(paragraph, bookmark, bookmark_id)
    return paragraph


def add_bullet(doc: Document, text: str, num_id: int):
    paragraph = doc.add_paragraph(style="Phaeno Bullet")
    numpr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    numpr.get_or_add_ilvl().val = 0
    numpr.get_or_add_numId().val = num_id
    run = paragraph.add_run(text)
    set_run_font(run, size=9.75, color=NEUTRAL_900)
    return paragraph


def add_numbered(doc: Document, text: str, num_id: int):
    paragraph = doc.add_paragraph(style="Phaeno Numbered")
    numpr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    numpr.get_or_add_ilvl().val = 0
    numpr.get_or_add_numId().val = num_id
    run = paragraph.add_run(text)
    set_run_font(run, size=9.75, color=NEUTRAL_900)
    return paragraph


def create_numbering(doc: Document, *, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "279")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.extend([tabs, ind, spacing])
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), GREEN_700 if kind == "bullet" else RNA_600)
    rpr.extend([fonts, color])
    level.extend([start, fmt, text, jc, ppr, rpr])
    abstract.append(level)
    # Word requires every abstractNum before the first concrete num instance.
    first_num_index = next(
        (index for index, child in enumerate(numbering) if child.tag == qn("w:num")),
        len(numbering),
    )
    numbering.insert(first_num_index, abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_toc(doc: Document, entries: list[tuple[int, str, str, int]]) -> None:
    first_paragraph = None
    last_paragraph = None
    for index, (level, title, bookmark, page_number) in enumerate(entries):
        paragraph = doc.add_paragraph(style=f"TOC {level}")
        if index == 0:
            first_paragraph = paragraph
            begin_run = paragraph.add_run()
            begin = OxmlElement("w:fldChar")
            begin.set(qn("w:fldCharType"), "begin")
            begin.set(qn("w:dirty"), "true")
            begin_run._r.append(begin)
            instruction_run = paragraph.add_run()
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = ' TOC \\o "1-2" \\h \\z \\u '
            instruction_run._r.append(instr)
            separate_run = paragraph.add_run()
            separate = OxmlElement("w:fldChar")
            separate.set(qn("w:fldCharType"), "separate")
            separate_run._r.append(separate)

        add_hyperlink(paragraph, title, bookmark)
        tab_run = paragraph.add_run("\t")
        set_run_font(tab_run, size=10.25 if level == 1 else 9.5, color=NEUTRAL_900)
        page_run = paragraph.add_run(str(page_number))
        set_run_font(page_run, size=10.25 if level == 1 else 9.5, color=NEUTRAL_600)
        last_paragraph = paragraph

    if first_paragraph is None or last_paragraph is None:
        raise ValueError("TOC entries are required")
    end_run = last_paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size=9.75, color=NEUTRAL_900)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    set_widow_control(normal)

    h1 = styles["Heading 1"]
    set_style_font(h1, size=17, color=RNA_900, bold=True)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing = 1.15
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True

    h2 = styles["Heading 2"]
    set_style_font(h2, size=12.5, color=RNA_600, bold=True)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(5)
    h2.paragraph_format.line_spacing = 1.15
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True

    h3 = styles["Heading 3"]
    set_style_font(h3, size=10.5, color=GREEN_700, bold=True)
    h3.paragraph_format.space_before = Pt(7)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.line_spacing = 1.15
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.keep_together = True

    caption = styles["Caption"]
    set_style_font(caption, size=9.25, color=NEUTRAL_600, italic=True)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.line_spacing = 1.15
    caption.paragraph_format.keep_with_next = True

    for name, size, color, bold, italic in [
        ("Phaeno Cover Kicker", 9.5, GREEN_700, True, False),
        ("Phaeno Cover Title", 28, RNA_900, True, False),
        ("Phaeno Cover Subtitle", 13, RNA_600, False, False),
        ("Phaeno Metadata", 9.5, NEUTRAL_600, False, False),
        ("Phaeno Contents Title", 24, RNA_900, True, False),
        ("Phaeno Lead", 11.5, RNA_900, False, False),
        ("Phaeno Callout", 9.5, NEUTRAL_900, False, False),
        ("Phaeno Instruction", 9, NEUTRAL_600, False, True),
        ("Phaeno Citation", 8.5, NEUTRAL_600, False, False),
        ("Phaeno Bullet", 9.75, NEUTRAL_900, False, False),
        ("Phaeno Numbered", 9.75, NEUTRAL_900, False, False),
    ]:
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(style, size=size, color=color, bold=bold, italic=italic)

    styles["Phaeno Cover Kicker"].paragraph_format.space_after = Pt(12)
    styles["Phaeno Cover Kicker"].paragraph_format.keep_with_next = True
    styles["Phaeno Cover Title"].paragraph_format.space_after = Pt(8)
    styles["Phaeno Cover Title"].paragraph_format.line_spacing = 1.04
    styles["Phaeno Cover Title"].paragraph_format.keep_with_next = True
    styles["Phaeno Cover Subtitle"].paragraph_format.space_after = Pt(20)
    styles["Phaeno Cover Subtitle"].paragraph_format.line_spacing = 1.25
    styles["Phaeno Metadata"].paragraph_format.space_after = Pt(3)
    styles["Phaeno Contents Title"].paragraph_format.space_after = Pt(6)
    styles["Phaeno Contents Title"].paragraph_format.keep_with_next = True
    styles["Phaeno Lead"].paragraph_format.space_after = Pt(14)
    styles["Phaeno Lead"].paragraph_format.line_spacing = 1.3
    styles["Phaeno Callout"].paragraph_format.space_before = Pt(5)
    styles["Phaeno Callout"].paragraph_format.space_after = Pt(12)
    styles["Phaeno Callout"].paragraph_format.line_spacing = 1.25
    styles["Phaeno Callout"].paragraph_format.left_indent = Twips(180)
    styles["Phaeno Callout"].paragraph_format.right_indent = Twips(180)
    styles["Phaeno Instruction"].paragraph_format.space_after = Pt(14)
    styles["Phaeno Instruction"].paragraph_format.line_spacing = 1.2
    styles["Phaeno Citation"].paragraph_format.space_before = Pt(4)
    styles["Phaeno Citation"].paragraph_format.space_after = Pt(4)
    styles["Phaeno Citation"].paragraph_format.line_spacing = 1.15
    for list_name in ("Phaeno Bullet", "Phaeno Numbered"):
        styles[list_name].paragraph_format.left_indent = Twips(540)
        styles[list_name].paragraph_format.first_line_indent = Twips(-279)
        styles[list_name].paragraph_format.space_after = Pt(4)
        styles[list_name].paragraph_format.line_spacing = 1.208

    for name, level in (("TOC 1", 1), ("TOC 2", 2)):
        try:
            style = styles[name]
        except KeyError:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        set_style_font(
            style,
            size=10.25 if level == 1 else 9.5,
            color=NEUTRAL_900 if level == 1 else NEUTRAL_600,
            bold=level == 1,
        )
        style.paragraph_format.left_indent = Inches(0 if level == 1 else 0.28)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(7 if level == 1 else 4)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.insert(0, update)
    update.set(qn("w:val"), "true")
    even_odd = settings.find(qn("w:evenAndOddHeaders"))
    if even_odd is not None:
        settings.remove(even_odd)

    doc.core_properties.title = "Phaeno White Paper Template"
    doc.core_properties.subject = "Reusable scientific white paper template"
    doc.core_properties.author = "Phaeno"
    doc.core_properties.keywords = "Phaeno, white paper, scientific publication, template"
    doc.core_properties.comments = "Built from the Phaeno Website design system."


def configure_page_geometry(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def set_section_columns(section, count: int, *, space_dxa: int = 432) -> None:
    cols = section._sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        section._sectPr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_dxa if count > 1 else 0))
    cols.set(qn("w:equalWidth"), "1")


def add_section(doc: Document, start_type: WD_SECTION, *, columns: int):
    section = doc.add_section(start_type)
    configure_page_geometry(section)
    set_section_columns(section, columns)
    # python-docx copies the prior section properties, including explicit page
    # restarts. New scientific-layout sections should continue numbering.
    page_numbering = section._sectPr.find(qn("w:pgNumType"))
    if page_numbering is not None:
        section._sectPr.remove(page_numbering)
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True
    return section


def add_column_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_break(WD_BREAK.COLUMN)


def clear_story(story) -> None:
    paragraph = story.paragraphs[0]
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def configure_cover_footer(section) -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    clear_story(footer)
    paragraph = footer.paragraphs[0]
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = paragraph.add_run("WHITE PAPER")
    set_run_font(left, size=8, color=NEUTRAL_500, bold=True, all_caps=True, tracking_twips=18)
    paragraph.add_run("\t")
    right = paragraph.add_run("phaenobiotech.com")
    set_run_font(right, size=8, color=NEUTRAL_500)


def configure_body_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_story(section.header)
    clear_story(section.footer)

    header = section.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    header.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    label = header.add_run("WHITE PAPER")
    set_run_font(label, size=8, color=GREEN_700, bold=True, all_caps=True, tracking_twips=18)
    header.add_run("\t")
    title = header.add_run("[SHORT TITLE]")
    set_run_font(title, size=8, color=NEUTRAL_500)

    footer = section.footer.paragraphs[0]
    footer.paragraph_format.space_before = Pt(0)
    footer.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    site = footer.add_run("phaenobiotech.com")
    set_run_font(site, size=8, color=NEUTRAL_500)
    footer.add_run("\t")
    page_label = footer.add_run("Page ")
    set_run_font(page_label, size=8, color=NEUTRAL_500)
    add_field(footer, "PAGE", "1")


def restart_page_numbering(section, start: int = 1) -> None:
    sectpr = section._sectPr
    pgn = sectpr.find(qn("w:pgNumType"))
    if pgn is None:
        pgn = OxmlElement("w:pgNumType")
        sectpr.append(pgn)
    pgn.set(qn("w:start"), str(start))


def make_cover_crop() -> Path:
    WORK.mkdir(parents=True, exist_ok=True)
    out = WORK / "cover-molecular-header.png"
    with Image.open(COVER_IMAGE) as source:
        image = source.convert("RGBA")
        width, height = image.size
        crop_height = min(height, round(width / 1.92))
        top = max(0, (height - crop_height) // 2)
        cropped = image.crop((0, top, width, top + crop_height))
        scrim = Image.new("RGBA", cropped.size, (14, 40, 65, 74))
        cropped = Image.alpha_composite(cropped, scrim)
        with Image.open(LOGO) as logo_source:
            logo = logo_source.convert("RGBA")
            logo_width = 250
            logo_height = round(logo.height * logo_width / logo.width)
            logo = logo.resize((logo_width, logo_height), Image.Resampling.LANCZOS)
            cropped.alpha_composite(logo, (70, 52))
        draw = ImageDraw.Draw(cropped)
        draw.rectangle((0, cropped.height - 12, cropped.width, cropped.height), fill="#789946")
        draw.rectangle((0, cropped.height - 12, 170, cropped.height), fill="#FEC950")
        cropped.convert("RGB").save(out, quality=94)
    return out


def make_figure_placeholder() -> Path:
    WORK.mkdir(parents=True, exist_ok=True)
    out = WORK / "evidence-figure-placeholder.png"
    width, height = 1400, 500
    image = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((2, 2, width - 3, height - 3), radius=22, outline="#D9D9D9", width=4, fill="#F7F8F8")
    draw.rectangle((0, 0, 20, height), fill="#789946")
    draw.rectangle((20, 0, width - 1, 14), fill="#0E2841")
    for y in (145, 230, 315, 400):
        draw.line((120, y, width - 95, y), fill="#D9D9D9", width=2)
    draw.line((120, 400, width - 95, 400), fill="#7F7F7F", width=3)
    draw.line((120, 115, 120, 400), fill="#7F7F7F", width=3)
    points = [(145, 355), (360, 320), (575, 335), (790, 250), (1005, 210), (1220, 145)]
    draw.line(points, fill="#789946", width=10, joint="curve")
    for x, y in points:
        draw.ellipse((x - 9, y - 9, x + 9, y + 9), fill="#FEC950", outline="#0E2841", width=2)
    regular = ImageFont.truetype(r"C:\Windows\Fonts\segoeui.ttf", 25)
    semibold = ImageFont.truetype(r"C:\Windows\Fonts\seguisb.ttf", 31)
    footnote = ImageFont.truetype(r"C:\Windows\Fonts\segoeui.ttf", 21)
    draw.text((125, 34), "FIGURE PLACEHOLDER", fill="#526832", font=semibold)
    draw.text((125, 76), "Replace with a conclusion-led chart, sequencing schematic, or microscopy image.", fill="#1D1D1D", font=regular)
    draw.text((125, 438), "Include units, n, method, source, and validation status. Illustration only.", fill="#595959", font=footnote)
    image.save(out)
    return out


def add_body_paragraph(doc: Document, text: str, *, italic_terms: list[str] | None = None):
    paragraph = doc.add_paragraph()
    if not italic_terms:
        run = paragraph.add_run(text)
        set_run_font(run, size=9.75, color=NEUTRAL_900)
        return paragraph
    cursor = 0
    for term in italic_terms:
        start = text.find(term, cursor)
        if start < 0:
            continue
        if start > cursor:
            run = paragraph.add_run(text[cursor:start])
            set_run_font(run, size=9.75, color=NEUTRAL_900)
        run = paragraph.add_run(term)
        set_run_font(run, size=9.75, color=NEUTRAL_900, italic=True)
        cursor = start + len(term)
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=9.75, color=NEUTRAL_900)
    return paragraph


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    cover_crop = make_cover_crop()
    figure_placeholder = make_figure_placeholder()

    doc = Document()
    set_document_defaults(doc)
    configure_styles(doc)
    bullet_num_id = create_numbering(doc, kind="bullet")
    sequence_num_id = create_numbering(doc, kind="decimal")
    recommended_num_id = create_numbering(doc, kind="decimal")
    reference_num_id = create_numbering(doc, kind="decimal")

    # Title page: editorial-cover pattern with a one-third-page discovery header.
    cover_section = doc.sections[0]
    configure_cover_footer(cover_section)

    image_paragraph = doc.add_paragraph()
    image_paragraph.paragraph_format.space_after = Pt(18)
    image_shape = image_paragraph.add_run().add_picture(str(cover_crop), width=Inches(6.5))
    set_picture_alt_text(image_shape, "Phaeno wordmark over an abstract molecular image")

    kicker = doc.add_paragraph(style="Phaeno Cover Kicker")
    kicker_run = kicker.add_run("WHITE PAPER TEMPLATE")
    set_run_font(kicker_run, size=9.5, color=GREEN_700, bold=True, all_caps=True, tracking_twips=28)

    title = doc.add_paragraph(style="Phaeno Cover Title")
    title_run = title.add_run("[White paper title]")
    set_run_font(title_run, size=28, color=RNA_900, bold=True)

    subtitle = doc.add_paragraph(style="Phaeno Cover Subtitle")
    subtitle_run = subtitle.add_run("[State the scientific conclusion or question in one clear sentence]")
    set_run_font(subtitle_run, size=13, color=RNA_600)

    meta = doc.add_paragraph(style="Phaeno Metadata")
    meta_run = meta.add_run("Prepared by [author or team]")
    set_run_font(meta_run, size=9.5, color=NEUTRAL_600, bold=True)
    meta2 = doc.add_paragraph(style="Phaeno Metadata")
    meta2_run = meta2.add_run("[Month Year]   |   Version [x.x]")
    set_run_font(meta2_run, size=9.5, color=NEUTRAL_600)
    meta2.paragraph_format.space_after = Pt(0)

    # Main document section begins at page 1; the cover remains unnumbered.
    body_section = add_section(doc, WD_SECTION.NEW_PAGE, columns=1)
    restart_page_numbering(body_section, 1)
    configure_body_header_footer(body_section)

    contents = doc.add_paragraph(style="Phaeno Contents Title")
    contents.add_run("Contents")
    note = doc.add_paragraph(style="Phaeno Instruction")
    note_run = note.add_run("Template note: after editing headings, press Ctrl+A, then F9 in Word to refresh the contents and page numbers.")
    set_run_font(note_run, size=9, color=NEUTRAL_600, italic=True)
    set_paragraph_shading(note, NEUTRAL_50)
    set_paragraph_left_border(note, NEUTRAL_300, size=10, space=6)

    toc_entries = [
        (1, "Executive summary", "exec_summary", 2),
        (2, "At a glance", "at_a_glance", 2),
        (1, "Introduction", "introduction", 3),
        (2, "Background", "background", 3),
        (2, "Scope and definitions", "scope", 3),
        (2, "Approach", "approach", 3),
        (1, "Evidence and findings", "evidence_findings", 4),
        (2, "Finding 1: [Conclusion-led heading]", "finding_1", 4),
        (2, "Finding 2: [Conclusion-led heading]", "finding_2", 4),
        (1, "Interpretation and recommendations", "interpretation", 5),
        (2, "What the evidence means", "meaning", 5),
        (2, "Recommended actions", "recommendations", 5),
        (1, "Conclusion", "conclusion", 5),
        (1, "References", "references", 6),
        (2, "Acknowledgments and disclosures", "acknowledgments", 6),
        (1, "Appendix A. Methods and supporting detail", "appendix_a", 7),
        (2, "Method summary", "method_summary", 7),
        (2, "Supplemental data", "supplemental_data", 7),
    ]
    add_toc(doc, toc_entries)

    # Page 2: Executive summary in the standard two-column scientific grid.
    add_section(doc, WD_SECTION.NEW_PAGE, columns=2)
    add_heading(doc, "Executive summary", 1, "exec_summary", 1)
    lead = doc.add_paragraph(style="Phaeno Lead")
    lead_run = lead.add_run("[Lead with the answer. In two or three sentences, state what was studied, what the evidence shows, and why it matters.]")
    set_run_font(lead_run, size=11.5, color=RNA_900)
    add_callout(
        doc,
        "Key takeaway",
        "[Write the single finding or implication a reader should remember after closing the paper.]",
    )

    add_heading(doc, "At a glance", 2, "at_a_glance", 2)
    add_bullet(doc, "Problem — [Name the scientific or operational limitation.]", bullet_num_id)
    add_bullet(doc, "Evidence — [Name the dataset, method, and strongest observed result.]", bullet_num_id)
    add_bullet(doc, "Implication — [Explain the decision, workflow, or research consequence.]", bullet_num_id)
    add_bullet(doc, "Validation status — [Measured, inferred, preliminary, replicated, or validated.]", bullet_num_id)

    add_column_break(doc)
    add_heading(doc, "Evidence checklist", 3, "evidence_checklist", 19)
    add_bullet(doc, "[State the sample size, specimen type, and exclusions.]", bullet_num_id)
    add_bullet(doc, "[Name the assay, analysis workflow, software, and version.]", bullet_num_id)
    add_bullet(doc, "[Report units, uncertainty, controls, and acceptance criteria.]", bullet_num_id)
    add_bullet(doc, "[Separate measured, inferred, preliminary, and validated claims.]", bullet_num_id)

    add_heading(doc, "Drafting rule", 3, "drafting_rule", 20)
    add_body_paragraph(
        doc,
        "[Use the executive summary to make the result legible, not to repeat the full paper. Keep the strongest evidence and its qualification together.]",
    )
    add_callout(
        doc,
        "Evidence note",
        "Keep sample size, method, units, source, and validation status adjacent to every consequential claim.",
        kind="evidence",
    )

    # Page 3: Introduction.
    add_section(doc, WD_SECTION.NEW_PAGE, columns=2)
    add_heading(doc, "Introduction", 1, "introduction", 3)
    add_heading(doc, "Background", 2, "background", 4)
    add_body_paragraph(
        doc,
        "[Introduce the biological or technical context. Define the unresolved problem before presenting the product, platform, or result. Use concrete biological nouns and measurable verbs.]",
    )
    add_body_paragraph(
        doc,
        "[Summarize the state of current evidence. Distinguish established knowledge from preliminary observations, and cite the primary source for each material claim.]",
    )

    add_heading(doc, "Scope and definitions", 2, "scope", 5)
    add_body_paragraph(
        doc,
        "[State what the paper covers, what it does not cover, the intended audience, and any research-use-only or validation boundaries.]",
    )
    add_bullet(doc, "Population or samples — [Cohort, specimen type, inclusion criteria, exclusions.]", bullet_num_id)
    add_bullet(doc, "Method — [Assay, workflow, platform, software, and version.]", bullet_num_id)
    add_bullet(doc, "Endpoints — [Primary measurements and acceptance criteria.]", bullet_num_id)
    add_bullet(doc, "Terminology — [Define each acronym at first use.]", bullet_num_id)

    add_heading(doc, "Approach", 2, "approach", 6)
    add_body_paragraph(
        doc,
        "[Describe the approach at the level needed to evaluate the evidence. Put detailed protocols, parameter lists, and supplemental analyses in the appendix.]",
    )
    add_column_break(doc)
    add_heading(doc, "Suggested narrative sequence", 3, "narrative_sequence", 21)
    add_numbered(doc, "[Define the unresolved biological or technical problem.]", sequence_num_id)
    add_numbered(doc, "[Explain the mechanism or method used to address it.]", sequence_num_id)
    add_numbered(doc, "[Present evidence with the strongest result first.]", sequence_num_id)
    add_numbered(doc, "[Interpret the result and state its validation boundary.]", sequence_num_id)
    add_heading(doc, "Scientific voice", 3, "scientific_voice", 22)
    add_body_paragraph(
        doc,
        "[Prefer direct, exact language. Define acronyms on first use, avoid unsupported superlatives, and state limitations where the claim appears.]",
    )

    # Page 4: Evidence and findings. A full-width figure interrupts the grid only
    # where it materially improves scientific reading.
    add_section(doc, WD_SECTION.NEW_PAGE, columns=1)
    add_heading(doc, "Evidence and findings", 1, "evidence_findings", 7)
    fig_caption = doc.add_paragraph(style="Caption")
    fig_label = fig_caption.add_run("Figure 1. ")
    set_run_font(fig_label, size=9.25, color=RNA_900, bold=True)
    fig_title = fig_caption.add_run("[Conclusion-led figure title]")
    set_run_font(fig_title, size=9.25, color=NEUTRAL_600, italic=True)
    fig_paragraph = doc.add_paragraph()
    fig_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_paragraph.paragraph_format.space_after = Pt(3)
    fig_shape = fig_paragraph.add_run().add_picture(str(figure_placeholder), width=Inches(6.3))
    set_picture_alt_text(fig_shape, "Illustrative evidence figure placeholder with a green line and amber data markers")
    fig_source = doc.add_paragraph(style="Phaeno Citation")
    fig_source.add_run("Source: [Dataset, method, sample size, units, version, and validation status].")

    add_section(doc, WD_SECTION.CONTINUOUS, columns=2)
    add_heading(doc, "Finding 1: [Conclusion-led heading]", 2, "finding_1", 8)
    add_body_paragraph(
        doc,
        "[State the result first. Then explain the comparison, sample, uncertainty, and qualification that support it.]",
    )

    add_heading(doc, "Finding 2: [Conclusion-led heading]", 2, "finding_2", 9)
    add_body_paragraph(
        doc,
        "[Use a short paragraph for the second result. If the comparison requires more than a few values, move it into the full-width evidence table below.]",
    )

    add_section(doc, WD_SECTION.CONTINUOUS, columns=1)
    comparison = doc.add_table(rows=1, cols=3)
    headers = ["Evidence question", "Comparator or baseline", "Observed result"]
    for index, text in enumerate(headers):
        comparison.rows[0].cells[index].text = text
    for row_values in [
        ("[Metric or observation]", "[Reference method or condition]", "[Result with units and uncertainty]"),
        ("[Biological interpretation]", "[Current explanation]", "[Supported interpretation and caveat]"),
    ]:
        cells = comparison.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
    style_table(comparison, [2200, 3300, 3860], emphasize_column=2)
    table_source = doc.add_paragraph(style="Phaeno Citation")
    table_source.add_run("Table 1 source: [Citation or internal dataset identifier].")

    # Page 5: Interpretation, recommendations, and conclusion.
    add_section(doc, WD_SECTION.NEW_PAGE, columns=2)
    add_heading(doc, "Interpretation and recommendations", 1, "interpretation", 10)
    add_heading(doc, "What the evidence means", 2, "meaning", 11)
    add_body_paragraph(
        doc,
        "[Explain the result without overstating it. Separate direct measurement from inference, and state which alternative explanations remain plausible.]",
    )
    add_callout(
        doc,
        "Scientific interpretation",
        "[Connect the measured result to the biological or workflow implication. Keep any limitation in the same paragraph.]",
    )

    add_column_break(doc)
    add_heading(doc, "Recommended actions", 2, "recommendations", 12)
    add_numbered(doc, "[Action 1: name the next decision, experiment, or workflow change.]", recommended_num_id)
    add_numbered(doc, "[Action 2: name the owner, evidence threshold, or review point.]", recommended_num_id)
    add_numbered(doc, "[Action 3: name the validation or commercialization boundary.]", recommended_num_id)

    add_heading(doc, "Conclusion", 1, "conclusion", 13)
    add_body_paragraph(
        doc,
        "[Restate the central result in plain language, explain its practical significance, and identify the next evidence needed. Do not introduce a new claim here.]",
    )

    # Page 6: References.
    add_section(doc, WD_SECTION.NEW_PAGE, columns=2)
    add_heading(doc, "References", 1, "references", 14)
    reference_note = doc.add_paragraph(style="Phaeno Instruction")
    reference_note.add_run("Use one citation style consistently. Prefer primary sources and include persistent identifiers when available.")
    set_paragraph_shading(reference_note, NEUTRAL_50)
    set_paragraph_left_border(reference_note, RNA_600, size=10, space=6)
    add_numbered(doc, "[Author(s). Title. Journal or publisher. Year; volume(issue):pages. DOI or URL.]", reference_num_id)
    add_numbered(doc, "[Author(s). Dataset or protocol title. Repository. Version. Year. DOI or accession.]", reference_num_id)
    add_numbered(doc, "[Phaeno internal study or technical note identifier. Version. Review date.]", reference_num_id)

    add_column_break(doc)
    add_heading(doc, "Acknowledgments and disclosures", 2, "acknowledgments", 15)
    add_body_paragraph(
        doc,
        "[List contributors, funding, conflicts of interest, data availability, code availability, and applicable research-use-only statements.]",
    )
    add_heading(doc, "Publication details", 3, "publication_details", 23)
    add_body_paragraph(
        doc,
        "[Add the document identifier, version, review date, corresponding author, and a stable public URL. State whether the paper is research use only.]",
    )

    # Page 7: Appendix returns to one column for wider method and supplemental tables.
    add_section(doc, WD_SECTION.NEW_PAGE, columns=1)
    add_heading(doc, "Appendix A. Methods and supporting detail", 1, "appendix_a", 16)
    appendix_lead = doc.add_paragraph(style="Phaeno Lead")
    appendix_lead.add_run("[Place detail here that supports reproducibility or review but would interrupt the main argument.]")

    add_heading(doc, "Method summary", 2, "method_summary", 17)
    method_table = doc.add_table(rows=1, cols=3)
    for index, text in enumerate(("Parameter", "Required detail", "Template prompt")):
        method_table.rows[0].cells[index].text = text
    for row_values in [
        ("Samples", "Source, count, quality, exclusions", "[Describe cohort or specimens]"),
        ("Wet-lab workflow", "Protocol, kit, lot, instruments", "[Add versioned method details]"),
        ("Analysis", "Pipeline, parameters, references", "[Add software and thresholds]"),
        ("Statistics", "Model, uncertainty, multiplicity", "[Add analysis plan]"),
    ]:
        cells = method_table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = value
    style_table(method_table, [1900, 3700, 3760])
    method_source = doc.add_paragraph(style="Phaeno Citation")
    method_source.add_run("Table A1. Method details required for reproducibility and review.")

    add_heading(doc, "Supplemental data", 2, "supplemental_data", 18)
    add_bullet(doc, "[Supplemental figure, table, or sensitivity analysis.]", bullet_num_id)
    add_bullet(doc, "[Data and code availability statement.]", bullet_num_id)
    add_bullet(doc, "[Known limitation, edge case, or validation dependency.]", bullet_num_id)

    add_callout(
        doc,
        "Appendix note",
        "Add Appendix B, C, and so on only when the supporting material has a distinct purpose. Keep each appendix in the contents.",
        kind="evidence",
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
